from docx import Document
import os
from tkinter import filedialog

def getDiretorio():
    file = filedialog.askdirectory(mustexist=True)
    return(file)

def preencherTriagem(entradaDocx, saidaDocx, dados):
    
    if not os.path.exists(entradaDocx):
        raise FileNotFoundError(f"Arquivo não encontrado: {entradaDocx}")

    doc = Document(entradaDocx)
    
    tabela = doc.tables[2]

    tabela.cell(0, 0).text = "Data: " + dados.get("data","")
    tabela.cell(0, 1).text = "NF: "+ dados.get("NF", "")

    tabela.cell(0, 2).text = "Ticket: " + dados.get("ticket", "")
    tabela.cell(1, 1).text = "Técnico: "+ dados.get("tecnico", "")

    tabela.cell(1, 2).text = "Analista Responsável: "+dados.get("analistaResponsavel", "")
    tabela.cell(2, 1).text = "Região: "+dados.get("regiao", "")

    tabela.cell(2, 2).text = "Equipamento: "+dados.get("equipamento", "")
    tabela.cell(3, 1).text = "N° de Série: "+dados.get("NSerie", "")

    tabela.cell(4, 1).text = "Peça/Equipamento Trocado: "+dados.get("pecaEquipamentoTrocado", "")

    tabela.cell(5, 1).text = "Defeito Alegado: "+dados.get("defeitoAlegado", "")

    tabela.cell(6, 1).text = "Constatação do Técnico: "+dados.get("constatacaoTecnico", "")

    tabela2 = doc.tables[4]

    tabela2.cell(0, 0).text = "Defeito constatado: " +dados.get("defeitoConstatado", "")
    tabela2.cell(1, 0).text = "Testes Realizados: " +dados.get("testesRealizados", "")
    tabela2.cell(2, 0).text = "Análise: " +dados.get("analise", "")
    tabela2.cell(3, 0).text = "Conclusão: " +dados.get("conclusao", "")


    caminhoFinal = os.path.join(getDiretorio(), saidaDocx)
    doc.save(caminhoFinal)

    print(f"Documento preenchido e salvo: {saidaDocx}")

if __name__ == "__main__":

    dados = {
        "data": "00/00/0000",
        "NF": "-",
        "ticket": "-",
        "tecnico": "-",
        "analistaResponsavel": "-",
        "regiao": "-",
        "equipamento": "-",
        "NSerie": "-",
        "pecaEquipamentoTrocado": "-",
        "defeitoAlegado": "-",
        "constatacaoTecnico": "-",
        "defeitoConstatado": "-",
        "testesRealizados": "-",
        "analise": "-",
        "conclusao": "-"
    }


    preencherTriagem(
        "TRIAGEM.docx",
        f"{dados.get('pecaEquipamentoTrocado','')}_{dados.get('ticket','')}_{dados.get('NSerie','')}.docx",
        dados
    )



